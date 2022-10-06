import docx
import os
from tkinter import *
from tkinter import ttk
from PIL import ImageTk, Image
from tkinter import filedialog
from docx.shared import Inches
import subprocess

doc = docx.Document()
doc.add_heading('Вывод документа для печати', 0)

root = Tk()
root.title("Печать документа")
root.geometry("700x550+300+100")

def openNewWin():
    root1 = Tk()
    columns = ("Items", "Values")
    Treeview = ttk.Treeview(root1, height=18, show="headings", columns=columns)  #

    Treeview.column("Items", width=200, anchor='center')
    Treeview.column("Values", width=200, anchor='center')

    Treeview.heading("Items", text="Items")
    Treeview.heading("Values", text="Values")

    Treeview.pack(side=LEFT, fill=BOTH)

    name = ['Item1', 'Item2', 'Item3']
    ipcode = ['Value1', 'Value2', 'Value3']
    for i in range(min(len(name), len(ipcode))):
        Treeview.insert('', i, values=(name[i], ipcode[i]))

    def set_cell_value(event):
        for item in Treeview.selection():
            column = Treeview.identify_column(event.x)
            row = Treeview.identify_row(event.y)
        cn = int(str(column).replace('#', ''))
        rn = int(str(row).replace('I', ''))
        entryedit = Text(root1, width=10 + (cn - 1) * 16, height=1)
        entryedit.place(x=16 + (cn - 1) * 130, y=6 + rn * 20)

        def saveedit():
            Treeview.set(item, column=column, value=entryedit.get(0.0, "end"))
            entryedit.destroy()
            okb.destroy()

        okb = ttk.Button(root1, text='OK', width=4, command=saveedit)
        okb.place(x=90 + (cn - 1) * 242, y=2 + rn * 20)

    def newrow():
        name.append('to be named')
        ipcode.append('value')
        Treeview.insert('', len(name) - 1, values=(name[len(name) - 1], ipcode[len(name) - 1]))
        Treeview.update()
        newb.place(x=120, y=(len(name) - 1) * 20 + 45)
        newb.update()

    Treeview.bind('<Double-1>', set_cell_value)
    newb = ttk.Button(root1, text='new item', width=20, command=newrow)
    newb.place(x=120, y=(len(name) - 1) * 20 + 45)

    for col in columns:
        Treeview.heading(col, text=col)

    def save_file():
        with open("file.docx", "w", newline=''):
            doc.add_heading('Таблица', level=1)
            table = doc.add_table(rows=1, cols=2)
            row1 = table.rows[0].cells
            row1[0].text = 'Items'
            row1[1].text = 'Values'
            for values in Treeview.get_children():
                row = Treeview.item(values)['values']
                row2 = table.add_row().cells
                row2[0].text = row[0]
                row2[1].text = row[1]
                print('save row:', row2[0].text, " and", row2[1].text)
            root1.destroy()

    btn11 = Button(root1,
                   text="Сохранить таблицу",
                   background="White",
                   foreground="Black",
                   padx="20",
                   pady="8",
                   font=("Times New Roman", 14, "bold"),
                   command=save_file)
    btn11.pack()
    root1.mainloop()


def click_button():
    doc.add_paragraph()
    doc.add_paragraph(name_entry.get())
    doc.save('file.docx')
    print("File save")

    def generate_pdf(doc_path, path):
        subprocess.call(['soffice',
                         '--convert-to',
                         'pdf',
                         '--outdir',
                         path,
                         doc_path])
        return doc_path

    generate_pdf("file.docx", "")
    # os.system("lpr file.pdf")#pechatat'
    print("Print file")


def openfilename():
    filename = filedialog.askopenfilename(title='"pen')
    return filename


def open_img():
    x = openfilename()
    img = Image.open(x)
    size = 250, 250
    img = img.resize(size, Image.LANCZOS)
    img = ImageTk.PhotoImage(img)
    panel = Label(root, image=img)
    panel.image = img
    panel.grid(row=4)

    path = os.path.realpath(x)
    print(path)
    doc.add_picture(path, height=Inches(5))

name_label1 = Label(text="Введите текст:")
name_label1.grid(row=4, column=4, sticky="w", pady=50, padx=50)

name_entry = Entry()
name_entry.grid(row=4, column=4, pady=50, padx=200)

btn = Button(root,
             text="Распечатать файл",
             background="White",
             foreground="Black",
             padx="20",
             pady="8",
             font=("Times New Roman", 14, "bold"),
             command=click_button)
btn1 = Button(root,
              text="Ввести данные в таблицу",
              background="White",
              foreground="Black",
              padx="20",
              pady="8",
              font=("Times New Roman", 14, "bold"),
              command=openNewWin)
btn2 = Button(root,
              text="Загрузить изображение",
              background="White",
              foreground="Black",
              padx="20",
              pady="8",
              font=("Times New Roman", 14, "bold"),
              command=open_img)

btn.place(x=235, y=450)
btn1.place(x=380, y=350)
btn2.place(x=50, y=350)

root.mainloop()